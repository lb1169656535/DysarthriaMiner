import csv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def csv_to_word(csv_path, word_path):
    # 创建新文档并设置页面
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # 定义带编号的样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Inches(0)  # 取消首行缩进

    try:
        with open(csv_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            next(reader)  # 跳过标题行
            
            # 添加参考文献标题
            doc.add_heading('参考文献列表', 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 生成带编号的引用
            for idx, row in enumerate(reader, 1):
                if len(row) < 2:
                    print(f"警告：第{idx}条数据缺失，已保留编号")
                citation = row[1] if len(row)>=2 else "（数据缺失）"
                
                # 创建带格式的段落
                p = doc.add_paragraph(style='Normal')
                p.add_run(f"[{idx}] ").bold = True    # 加粗编号
                p.add_run(citation)                  # 正常字体内容
                p.paragraph_format.space_after = Pt(6)  # 段后间距

    except Exception as e:
        print(f"处理失败: {str(e)}")
        return

    doc.save(word_path)
    print(f"成功生成包含 {idx} 条编号引用的文档：{word_path}")

# 使用示例
if __name__ == "__main__":
    csv_to_word('citations.csv', 'references.docx')
