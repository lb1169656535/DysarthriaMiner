import csv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def process_csv(input_csv, output_docx):
    entries = {}
    duplicates = []

    with open(input_csv, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            key = (row['Title'].strip(), row['Abstract'].strip())
            if key in entries:
                duplicates.append(key[0])
                continue
            entries[key] = row

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 主标题
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.add_run('IEEE上关于dysarthria论文题目与摘要汇总')
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    # 条目格式控制
    for i, ((title, abstract), _) in enumerate(entries.items()):
        # 添加题目
        title_para = doc.add_paragraph()
        run = title_para.add_run(title)
        run.bold = True
        
        # 题目与摘要间两个换行
        doc.add_paragraph()  # 第一个换行
        #doc.add_paragraph()  # 第二个换行

        # 添加摘要
        abstract_para = doc.add_paragraph(abstract)
        abstract_para.paragraph_format.line_spacing = 1.5

        # 条目间三个换行（最后一个条目不加）
        if i < len(entries)-1:
            doc.add_paragraph()  # 第一个换行
            doc.add_paragraph()  # 第二个换行 
            #doc.add_paragraph()  # 第三个换行

    doc.save(output_docx)

    if duplicates:
        print(f'发现 {len(duplicates)} 条重复条目:')
        for dup in duplicates:
            print(f'- {dup}')
    else:
        print('未发现重复条目')

# 使用示例
process_csv('IEEE_paper_metadata.csv', 'Papers_Summary.docx')